"""
Document parsing service for extracting text from various file formats.

Supported formats:
- TXT (plain text)
- PDF (using pypdf)
- DOCX (using python-docx)
"""

import os
import logging
from typing import Optional, Dict, Any
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Represents a parsed document with extracted text and metadata"""
    content: str
    metadata: Dict[str, Any]
    page_count: Optional[int] = None
    word_count: int = 0
    char_count: int = 0


class DocumentParser:
    """Service for parsing various document formats into plain text"""
    
    SUPPORTED_EXTENSIONS = {'txt', 'pdf', 'docx'}
    
    def __init__(self):
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are available"""
        self.has_pypdf = False
        self.has_docx = False
        
        try:
            import pypdf
            self.has_pypdf = True
        except ImportError:
            logger.warning("pypdf not installed. PDF parsing will not be available.")
        
        try:
            import docx
            self.has_docx = True
        except ImportError:
            logger.warning("python-docx not installed. DOCX parsing will not be available.")
    
    def is_supported(self, filename: str) -> bool:
        """Check if a file format is supported"""
        if not filename or '.' not in filename:
            return False
        ext = filename.rsplit('.', 1)[1].lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def get_file_extension(self, filename: str) -> str:
        """Get the file extension from a filename"""
        if not filename or '.' not in filename:
            return ''
        return filename.rsplit('.', 1)[1].lower()
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a document and extract its text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParsedDocument with extracted text and metadata
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = self.get_file_extension(file_path)
        
        if ext == 'txt':
            return self._parse_txt(file_path)
        elif ext == 'pdf':
            if not self.has_pypdf:
                raise ValueError("PDF parsing requires pypdf. Install it with: pip install pypdf")
            return self._parse_pdf(file_path)
        elif ext == 'docx':
            if not self.has_docx:
                raise ValueError("DOCX parsing requires python-docx. Install it with: pip install python-docx")
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _parse_txt(self, file_path: str) -> ParsedDocument:
        """Parse a plain text file"""
        logger.info(f"Parsing TXT file: {file_path}")
        
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        content = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            # Fallback: read as bytes and decode with errors='ignore'
            with open(file_path, 'rb') as f:
                content = f.read().decode('utf-8', errors='ignore')
            used_encoding = 'utf-8 (with errors ignored)'
        
        word_count = len(content.split())
        char_count = len(content)
        
        return ParsedDocument(
            content=content,
            metadata={
                'format': 'txt',
                'encoding': used_encoding,
                'file_size': os.path.getsize(file_path)
            },
            page_count=None,
            word_count=word_count,
            char_count=char_count
        )
    
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse a PDF file"""
        logger.info(f"Parsing PDF file: {file_path}")
        
        import pypdf
        
        content_parts = []
        page_count = 0
        
        try:
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                page_count = len(reader.pages)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            content_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                # Get metadata
                pdf_metadata = {}
                if reader.metadata:
                    for key in ['/Title', '/Author', '/Subject', '/Creator', '/Producer']:
                        if key in reader.metadata:
                            pdf_metadata[key.lstrip('/')] = reader.metadata[key]
        
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")
        
        content = '\n\n'.join(content_parts)
        word_count = len(content.split())
        char_count = len(content)
        
        return ParsedDocument(
            content=content,
            metadata={
                'format': 'pdf',
                'file_size': os.path.getsize(file_path),
                **pdf_metadata
            },
            page_count=page_count,
            word_count=word_count,
            char_count=char_count
        )
    
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse a DOCX file"""
        logger.info(f"Parsing DOCX file: {file_path}")
        
        import docx
        
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            logger.error(f"Error opening DOCX file: {e}")
            raise ValueError(f"Failed to open DOCX file: {e}")
        
        content_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)
        
        # Extract tables
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_content.append(cell_text)
                if row_content:
                    table_content.append(' | '.join(row_content))
            if table_content:
                content_parts.append('\n'.join(table_content))
        
        content = '\n\n'.join(content_parts)
        word_count = len(content.split())
        char_count = len(content)
        
        # Get core properties
        doc_metadata = {}
        try:
            core_props = doc.core_properties
            if core_props.title:
                doc_metadata['Title'] = core_props.title
            if core_props.author:
                doc_metadata['Author'] = core_props.author
            if core_props.subject:
                doc_metadata['Subject'] = core_props.subject
        except Exception as e:
            logger.warning(f"Could not extract DOCX metadata: {e}")
        
        return ParsedDocument(
            content=content,
            metadata={
                'format': 'docx',
                'file_size': os.path.getsize(file_path),
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                **doc_metadata
            },
            page_count=None,  # DOCX doesn't have a reliable page count
            word_count=word_count,
            char_count=char_count
        )
